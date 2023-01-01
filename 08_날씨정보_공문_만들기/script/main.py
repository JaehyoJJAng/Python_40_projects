import requests as rq
import os
import re
import docx
from docx import Document
from docx.oxml.ns import qn
from typing import Dict

class GetWeatherData:
    def __init__(self)-> None:
        self.url : str = 'http://www.kma.go.kr/wid/queryDFSRSS.jsp?zone=4159054000'

    def parsing(self)-> Dict[str,str]:
        """ 날씨 정보 파싱 및 리턴 """
        with rq.Session() as session:
            with session.get(url=self.url) as res:
                if res.ok:
                    weather_data : dict = dict()

                    # 요청 시간
                    pubDate = re.findall(r'<pubDate>(.+)</pubDate>',str(res.text))[0]

                    # 지역
                    area = re.findall(r'<category>(.+)</category>',res.text)[0]

                    # 날씨 정보
                    wfKor = re.findall(r'<wfKor>(.+)</wfKor>',res.text)

                    # 시간
                    hour = re.findall(r'<hour>(.+)</hour>',res.text)

                    # 온도
                    temp = re.findall(r'<temp>(.+)</temp>',res.text)

                    # 습도
                    reh = re.findall(r'<reh>(.+)</reh>',res.text)

                    weather_data['pubDate'] = pubDate
                    weather_data['area'] = area
                    weather_data['wfKor'] = wfKor
                    weather_data['hour'] = hour
                    weather_data['temp'] = temp
                    weather_data['reh'] = reh
        return weather_data

class Docx:
    def __init__(self,docxFile:str)-> None:
        self.docxFile = docxFile

    def load_docx(self)-> None:
        """ 워드 문서 읽어오기 """
        doc = Document(self.docxFile)

        # 워드문서 내용 출력
        for i , paragraph in enumerate(doc.paragraphs):
            print(f'{i} : {paragraph.text}')

        # 워드문서의 표 출력
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    print(para.text)

def main()-> None:
    # GetWeatherData Instance
    weather = GetWeatherData()

    # 날찌 정보 파싱
    weather_data : dict = weather.parsing()

    # docx 파일
    docxFile : str = '../docx/공문생성.docx'

    # Docx Instance
    doc = Docx(docxFile=docxFile)

    doc.load_docx()

if __name__ == '__main__':
    main()
