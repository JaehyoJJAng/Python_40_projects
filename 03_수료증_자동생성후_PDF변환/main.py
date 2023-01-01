"""
pip install python-docx
pip install docx2pdf
"""
from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn
from docx2pdf import convert
import docx
import os
from typing import List,Dict

class Xlsx:
    def __init__(self,xlsxPath:str)-> None:
        self.xlsxPath : str = xlsxPath 
    
    def load_data(self)-> List[Dict[str,str]]:

        """ 엑셀에서 데이터 읽어오는 메서드 """
        wb = load_workbook(os.path.join(self.xlsxPath,'수료명단.xlsx'))
        ws = wb.active
    
        datas = list()
        for i in range(2,ws.max_row + 1):
            data_dic = dict()
            
            # Dictionary Data
            data_dic['name'] = ws.cell(row=i,column=1).value
            data_dic['birthday'] = ws.cell(row=i,column=2).value
            data_dic['ho'] = ws.cell(row=i,column=3).value
            
            # Add in List
            datas.append(data_dic)
        return datas

class Word:
    def __init__(self,docxPath:str,datas:list)-> None:
        self.docxPath : str = docxPath
        self.datas : list = datas

    def load_docx(self)-> None:
        """ 워드파일 읽어오는 메서드 """
        doc = Document(os.path.join(self.docxPath,'수료증양식.docx'))
        
        text = ''
        for i , paragraph in enumerate(doc.paragraphs):
            text += str(i) + ": " + paragraph.text + '\n'
        print(text)

    def edit_word_file(self)-> None:
        """ 워드 파일 수정하는 메소드 """
        doc = Document(os.path.join(self.docxPath,'수료증양식.docx'))

        for data in self.datas:
            doc.paragraphs[3].clear()
            
            # 3번째 줄 수정
            run = doc.paragraphs[3].add_run('제 ' + data['ho'] +' 호')

            # Font Size
            run.font.size = docx.shared.Pt(20)
            
            # 6번째 줄 수정
            doc.paragraphs[6].clear()
            run = doc.paragraphs[6].add_run('성 명: ' + data['name'])

            # Font Size
            run.font.size = docx.shared.Pt(20)

            # 7번째 줄 수정
            doc.paragraphs[7].clear()
            run = doc.paragraphs[7].add_run('생 년 월 일: ' + data['birthday'])
            
            # 저장 경로 지정
            savePath = os.path.abspath('수료증')
            fileName = f'{data["ho"]}{data["name"]}.docx'
            if not os.path.exists(savePath):
                os.mkdir(savePath)
            
            # Word 저장
            doc.save(os.path.join(savePath,fileName))

            # PDF 변환
            convert(os.path.join(savePath,fileName),os.path.join(savePath,fileName.replace('.docx','.pdf'))) 

def main()-> None:
    xlsxPath = './xlsx'
    docxPath = './docx'

    datas : list = Xlsx(xlsxPath=xlsxPath).load_data()

    Word(docxPath=docxPath,datas=datas).edit_word_file()

if __name__ == '__main__':
    main()
